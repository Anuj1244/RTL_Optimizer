from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border:
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def create_report_docx(output_file="Project_Report.docx"):
    doc = Document()

    # --- Title Section ---
    title = doc.add_heading('AI-POWERED RTL OPTIMIZATION PROJECT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Introduction ---
    doc.add_heading('1. Project Objective', level=1)
    p = doc.add_paragraph(
        "The objective of this project is to automate the power optimization of RTL (Register-Transfer Level) designs. "
        "By identifying registers with high switching activity from power reports, the tool automatically injects "
        "sequential enable logic (register-level clock gating) using Large Language Models (LLMs) to reduce dynamic power consumption."
    )

    # --- Why LLM? ---
    doc.add_heading('2. Why We Use Large Language Models (LLM)', level=1)
    p = doc.add_paragraph()
    p.add_run("RTL designs are complex and vary greatly in coding styles. LLMs are used because:").bold = True
    bullets = [
        "Contextual Intelligence: Unlike regular expressions, LLMs understand the nested logic of 'always' blocks.",
        "Precision Refactoring: They can modify specific lines while ensuring the rest of the Verilog code remains intact.",
        "Adaptive Style: They handle different indentation and naming conventions seamlessly.",
        "Zero-Shot Optimization: They can apply logic like 'if (enable)' wraps without needing pre-defined templates for every case."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # --- Flow Diagram Section ---
    doc.add_heading('3. Project Flow Diagram', level=1)
    
    flow_steps = [
        "Input: Power Analysis Reports (.txt)",
        "↓",
        "Step 1: Parse Report (Extract Target Reg & Enable Signal)",
        "↓",
        "Step 2: RTL Context Extraction (Locate Always Block)",
        "↓",
        "Step 3: LLM Optimization (Llama 3.3 70B Refactoring)",
        "↓",
        "Step 4: Logic Injection & RTL Generation",
        "↓",
        "Step 5: Validation & Diff Reporting (Output: Optimized RTL)"
    ]

    table = doc.add_table(rows=len(flow_steps), cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, step in enumerate(flow_steps):
        cell = table.rows[i].cells[0]
        cell.text = step
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Step" in step or "Input" in step:
            cell.paragraphs[0].runs[0].bold = True
            # Optional: Add borders to make it look like boxes
            set_cell_border(cell, top={"sz": 12, "val": "single"}, bottom={"sz": 12, "val": "single"}, left={"sz": 12, "val": "single"}, right={"sz": 12, "val": "single"})

    # --- Step-by-Step Breakdown ---
    doc.add_heading('4. Detailed Optimization Steps', level=1)
    
    # Define steps with I/O
    steps_info = [
        ("1. Report Parsing", "Power Analysis Report (.txt)", "Metadata Dictionary (Register, Line, File, Enable)", "To identify the specific registers consuming excessive power."),
        ("2. Context Extraction", "Original RTL File + Metadata", "Verilog Always-Block Snippet", "The LLM needs context to ensure the modified code is syntactically correct."),
        ("3. LLM Refactoring", "Context + Enable Signal", "Optimized RTL Code Line", "The LLM intelligently wraps the assignment in an 'if' condition."),
        ("4. File Synthesis", "Original RTL + LLM Patch", "Optimized .v / .sv File", "To produce the final synthesizable hardware description."),
        ("5. Validation", "Optimized RTL File", "Success/Failure Status + Diff Report", "Ensures no syntax errors were introduced during the automated process.")
    ]

    for title_text, inp, out, why in steps_info:
        doc.add_heading(title_text, level=2)
        p = doc.add_paragraph()
        p.add_run("Input: ").bold = True
        p.add_run(inp)
        p = doc.add_paragraph()
        p.add_run("Output: ").bold = True
        p.add_run(out)
        p = doc.add_paragraph()
        p.add_run("Necessity: ").bold = True
        p.add_run(why)

    # --- Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "By combining traditional power analysis with cutting-edge AI (LLMs), this project demonstrates a modern "
        "approach to hardware design automation. It significantly reduces manual engineering effort and minimizes "
        "human error in the power optimization phase of chip design."
    )

    doc.save(output_file)
    print(f"Report saved as {output_file}")

if __name__ == "__main__":
    create_report_docx()
